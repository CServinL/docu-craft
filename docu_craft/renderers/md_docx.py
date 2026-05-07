"""Markdown → DOCX transformer via python-docx.

Emoji characters are left in the document and rely on Word's built-in font
fallback (Segoe UI Emoji on Windows, Apple Color Emoji on Mac) to render.
Stripping only happens when the caller explicitly passes strip_emoji=True.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from .base import BaseTransformer
from ..themes.base import resolve_style, _DEFAULT_STYLES

_PAGE_WIDTH_CM = 16.0

# Comprehensive emoji regex — used only when explicit stripping is requested
_EMOJI_RE = re.compile(
    "["
    "\U0001F000-\U0001FAFF"  # All extended emoji blocks
    "\U00002300-\U000027BF"  # Misc technical, symbols, dingbats
    "\U00002B00-\U00002BFF"  # Misc symbols and arrows
    "\U0000FE00-\U0000FE0F"  # Variation selectors (stray U+FE0F etc.)
    "\U0000200D"             # Zero-width joiner
    "\U000020E3"             # Combining enclosing keycap (1️⃣ etc.)
    "]+",
    re.UNICODE,
)

_PARA_STYLES = {
    "body":           "DC Body Text",
    "code_block":     "DC Code Block",
    "quote":          "DC Quote",
    "table_header":   "DC Table Header",
    "table_cell":     "DC Table Cell",
    "table_cell_alt": "DC Table Cell Alt",
}

_CHAR_STYLES = {
    "bold":        "DC Bold",
    "italic":      "DC Italic",
    "code_inline": "DC Code Inline",
}

# Built-in DOCX list styles support up to 3 levels natively
_UL_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_OL_STYLES = ["List Number", "List Number 2", "List Number 3"]


def _hex_to_rgb(hex_color: str) -> tuple:
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_bottom_border(cell, hex_color: str = "E0E0E0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        if side == "bottom":
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), hex_color.lstrip("#"))
        else:
            border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _list_item(line: str):
    """Return (indent, is_ordered, text) if line is a list item, else None."""
    m = re.match(r'^(\s*)([-*+])\s+(.*)', line)
    if m:
        return len(m.group(1)), False, m.group(3)
    m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
    if m:
        return len(m.group(1)), True, m.group(2)
    return None


class MdDocxTransformer(BaseTransformer):
    """Markdown → DOCX."""

    input_fmt = "md"
    output_fmt = "docx"
    applies_style = True
    priority = 1

    def transform(self, content: str, **options) -> bytes:
        self._style_cfg   = options.get("style", {})
        self._fonts       = self._style_cfg.get("fonts", {})
        self._styles      = self._style_cfg.get("styles", _DEFAULT_STYLES)
        self._colors      = self._style_cfg.get("colors", {})
        self._strip_emoji = options.get("strip_emoji", False)

        doc = Document()
        self._register_styles(doc)
        self._convert(doc, content)

        output = options.get("output")
        if output:
            doc.save(str(output))
            return Path(output)
        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _clean(self, text: str) -> str:
        if self._strip_emoji:
            return _EMOJI_RE.sub("", text).strip()
        return text

    # ------------------------------------------------------------------
    # Style registration
    # ------------------------------------------------------------------

    def _register_styles(self, doc):
        for level in range(1, 7):
            key = f"heading{level}"
            sd = resolve_style(self._styles.get(key, {}), self._fonts, "docx")
            try:
                h = doc.styles[f"Heading {level}"]
                h.font.name   = sd.get("font_name", "Arial")
                h.font.size   = Pt(sd.get("size", 12))
                h.font.bold   = sd.get("bold", True)
                h.font.italic = sd.get("italic", False)
                if sd.get("color"):
                    h.font.color.rgb = RGBColor(*_hex_to_rgb(sd["color"]))
                if "space_before" in sd:
                    h.paragraph_format.space_before = Pt(sd["space_before"])
                if "space_after" in sd:
                    h.paragraph_format.space_after  = Pt(sd["space_after"])
            except KeyError:
                pass

        for key, name in _PARA_STYLES.items():
            sd = resolve_style(self._styles.get(key, {}), self._fonts, "docx")
            self._add_para_style(doc, name, sd)

        for key, name in _CHAR_STYLES.items():
            sd = resolve_style(self._styles.get(key, {}), self._fonts, "docx")
            self._add_char_style(doc, name, sd)

        body_sd = resolve_style(self._styles.get("body", {}), self._fonts, "docx")
        try:
            doc.styles["Normal"].font.name = body_sd.get("font_name", "Calibri")
            doc.styles["Normal"].font.size = Pt(body_sd.get("size", 11))
        except KeyError:
            pass

    def _add_para_style(self, doc, name: str, sd: dict):
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = doc.styles[name]
        style.font.name = sd.get("font_name", "Calibri")
        if "size" in sd:
            style.font.size = Pt(sd["size"])
        if sd.get("bold"):
            style.font.bold = True
        if sd.get("italic"):
            style.font.italic = True
        if sd.get("color"):
            style.font.color.rgb = RGBColor(*_hex_to_rgb(sd["color"]))
        if "space_before" in sd:
            style.paragraph_format.space_before = Pt(sd["space_before"])
        if "space_after" in sd:
            style.paragraph_format.space_after  = Pt(sd["space_after"])

    def _add_char_style(self, doc, name: str, sd: dict):
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        except ValueError:
            style = doc.styles[name]
        style.font.name = sd.get("font_name", "Calibri")
        if "size" in sd:
            style.font.size = Pt(sd["size"])
        if sd.get("bold"):
            style.font.bold = True
        if sd.get("italic"):
            style.font.italic = True
        if sd.get("color"):
            style.font.color.rgb = RGBColor(*_hex_to_rgb(sd["color"]))

    # ------------------------------------------------------------------
    # Content conversion
    # ------------------------------------------------------------------

    def _convert(self, doc, text: str):
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Fenced code block
            if re.match(r"^```", line):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph(style=_PARA_STYLES["code_block"])
                p.add_run("\n".join(code_lines))
                i += 1
                continue

            # ATX headings
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                level = min(len(m.group(1)), 6)
                doc.add_heading(self._clean(_strip_inline(m.group(2))), level=level)
                i += 1
                continue

            # Table
            if re.match(r"^\|", line):
                rows = []
                while i < len(lines) and re.match(r"^\|", lines[i]):
                    if re.match(r"^\|[-| :]+\|", lines[i]):
                        i += 1
                        continue
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                    i += 1
                if rows:
                    self._add_table(doc, rows)
                continue

            # Horizontal rule
            if re.match(r"^[-*_]{3,}\s*$", line):
                doc.add_paragraph("─" * 50)
                i += 1
                continue

            # Blockquote
            if line.startswith("> "):
                p = doc.add_paragraph(style=_PARA_STYLES["quote"])
                _add_inline(p, line[2:], _CHAR_STYLES)
                i += 1
                continue

            # List — ordered or unordered, any indent level
            if _list_item(line) is not None:
                base_indent = _list_item(line)[0]
                while i < len(lines):
                    parsed = _list_item(lines[i])
                    if parsed is None:
                        break
                    indent, is_ordered, item_text = parsed
                    level = min(max((indent - base_indent) // 2, 0), 2)
                    style = (_OL_STYLES if is_ordered else _UL_STYLES)[level]
                    p = doc.add_paragraph(style=style)
                    _add_inline(p, item_text, _CHAR_STYLES)
                    i += 1
                continue

            # Blank line
            if line.strip() == "":
                i += 1
                continue

            # Normal paragraph
            p = doc.add_paragraph(style=_PARA_STYLES["body"])
            _add_inline(p, line, _CHAR_STYLES)
            i += 1

    def _add_table(self, doc, rows):
        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        col_width = Cm(_PAGE_WIDTH_CM / cols)
        for col in table.columns:
            col.width = col_width

        header_bg  = self._styles.get("table_header", {}).get("background", "#1a1a2e")
        header_clr = self._styles.get("table_header", {}).get("color", "#ffffff")
        alt_bg     = self._styles.get("table_cell_alt", {}).get("background", "#f7f7f9")
        border_clr = self._colors.get("border", "#e0e0e0")
        header_rgb = _hex_to_rgb(header_clr)

        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                cell = table.cell(r_idx, c_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                if r_idx == 0:
                    _set_cell_bg(cell, header_bg)
                    run = p.add_run(self._clean(_strip_inline(cell_text)))
                    run.style = doc.styles[_CHAR_STYLES["bold"]]
                    run.font.color.rgb = RGBColor(*header_rgb)
                else:
                    if r_idx % 2 == 0:
                        _set_cell_bg(cell, alt_bg)
                    _set_cell_bottom_border(cell, border_clr)
                    _add_inline(p, cell_text, _CHAR_STYLES)


# ------------------------------------------------------------------
# Inline helpers
# ------------------------------------------------------------------

def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text


def _add_inline(paragraph, text: str, char_styles: dict):
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))", re.DOTALL)
    for m in pattern.finditer(text):
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.style = paragraph.part.document.styles[char_styles["bold"]]
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.style = paragraph.part.document.styles[char_styles["italic"]]
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.style = paragraph.part.document.styles[char_styles["code_inline"]]
        elif m.group(5):
            if m.group(5):
                paragraph.add_run(m.group(5))
